import PyPDF2
import requests
import json
from docx import Document
import io
import os
from dotenv import load_dotenv
import logging
logging.basicConfig(level=logging.INFO, filename='app.log',format='%(asctime)s - %(levelname)s - %(message)s')

load_dotenv()
URL = os.getenv("URL")
API_KEY = os.getenv("API_KEY")
MODEL_NAME = os.getenv("MODEL_NAME")


class ResumeAnalyzer:
    def __init__(self, base_url=URL):
        self.base_url = base_url
        self.model_name = MODEL_NAME 

    def _call_llm_api(self, messages):
        """Make API call to the local LLM server."""
        try:
            #logging.info(f"ResumeAnalyzer: Message Recieved:{messages}")
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {API_KEY}"
            }
            
            payload = {
                "model": self.model_name,
                "messages": [{"role": "user","content":messages}]
            }
            
            response = requests.post(
                self.base_url,
                headers=headers,
                json=payload
            )
            
            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"]
            else:
                return f"API Error: {response.status_code}"
                
        except Exception as e:
            return f"Error calling LLM API: {str(e)}"
                
        except Exception as e:
            return f"Error calling Ollama API: {str(e)}"

    def extract_text_from_pdf(self, file):
        """Extract text from PDF file."""
        try:
            if hasattr(file, 'read'):
                # Handle file object
                pdf_reader = PyPDF2.PdfReader(file)
            else:
                # Handle file path
                with open(file, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            logging.info("ResumeAnalyzer: File identified as PDF")
            # logging.info({"From":["Start"],"Desc": "Extracting Text from PDF","To":"ResumeReader"}))
            return text
        except Exception as e:
            return f"Error extracting text from PDF: {str(e)}"

    def extract_text_from_docx(self, file):
        """Extract text from Word document."""
        try:
            if hasattr(file, 'read'):
                # Handle file object
                doc = Document(file)
            else:
                # Handle file path
                doc = Document(file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            #logging.info("ResumeAnalyzer: File identified as PDF"))
            logging.info({"From":["Start"],"Desc": "Extracting Text from docx","To":"ResumeReader"})
            return text
        except Exception as e:
            return f"Error extracting text from Word document: {str(e)}"

    def extract_text_from_file(self, file):
        """Extract text from a file (PDF or Word)."""
        try:
            # Handle both file objects and file paths
            if hasattr(file, 'name'):
                file_path = file.name
            else:
                file_path = file
                
            file_extension = file_path.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                return self.extract_text_from_pdf(file)
            elif file_extension == 'docx':
                return self.extract_text_from_docx(file)
            else:
                return f"Unsupported file format: {file_extension}"
                
        except Exception as e:
            return f"Error extracting text: {str(e)}"

    def analyze_resume(self, resume_text):
        """Analyze the resume and provide initial insights."""
        prompt = f"""
        Analyze the following resume and provide key insights:

        Resume:
        {resume_text}

        Please provide:
        1. Key Skills Summary:
           - Technical skills identified
           - Soft skills identified
           - Tools and technologies

        2. Experience Analysis:
           - Years of experience
           - Industry exposure
           - Key achievements
           - Leadership experience

        3. Education and Certifications:
           - Academic qualifications
           - Professional certifications
           - Additional training

        4. Strengths and Areas of Expertise:
           - Core competencies
           - Unique selling points
           - Notable accomplishments

        Format the response in a clear, structured way with sections and bullet points.
        """

        try:
            response = self._call_llm_api(prompt)
            logging.info({"From":["ResumeReader"],"Desc": f"""{response}""","To":"ResumeAnalyzer"})
            return response
        except Exception as e:
            return f"Error analyzing resume: {str(e)}"

    def get_skill_gap_analysis(self, resume_text, job_description):
        """Analyze the gap between resume skills and job requirements."""
        prompt = f"""
        Analyze the following resume and job description to identify skill gaps:
        
        Resume:
        {resume_text}
        
        Job Description:
        {job_description}
        
        Please provide:
        1. Missing technical skills required by the job
        2. Missing soft skills required by the job
        3. Experience gaps
        4. Specific recommendations to bridge these gaps
        
        Format the response in a clear, structured way with sections and bullet points.
        """
        
        try:
            response = self._call_llm_api(prompt)
            logging.info({"From":["ResumeAnalyzer","Start"],"Desc": f"""{response}""","To":"SkillGapAnalyzer"})
            return response
        except Exception as e:
            return f"Error analyzing skill gaps: {str(e)}"

    def get_career_path_advice(self, resume_text, job_description):
        """Provide advice on how to bridge the gap between current skills and job requirements."""
        prompt = f"""
        Based on the following resume and job description, provide a detailed career path plan:
        
        Resume:
        {resume_text}
        
        Job Description:
        {job_description}
        
        Please provide:
        1. Short-term actions (0-3 months)
        2. Medium-term goals (3-6 months)
        3. Long-term development plan (6-12 months)
        4. Recommended learning resources
        5. Potential certifications or courses
        6. Networking opportunities
        7. Project suggestions to build required skills
        
        Format the response in a clear, structured way with sections and bullet points.
        """
        
        try:
            response = self._call_llm_api(prompt)
            logging.info({"From":["Start"],"Desc": f"""{response.replace("'",'`').replace()}""","To":"CareerPathAdvisor"})
            return response
        except Exception as e:
            return f"Error generating career path advice: {str(e)}"

    def get_company_insights(self, company_name):
        """Generate insights about company interview process and culture."""
        prompt = f"""
        Generate detailed insights about {company_name}'s interview process and company culture.
        
        Please provide:
        1. Company Overview:
           - Company culture and values
           - Work environment
           - Employee benefits and perks
        
        2. Interview Process:
           - Typical interview stages
           - Timeline and duration
           - Common interview formats (remote/onsite)
        
        3. Interview Experience:
           - Common interview questions
           - Typical technical assessments
           - Behavioral interview focus areas
        
        4. Tips from Past Candidates:
           - Do's and don'ts
           - Common pitfalls
           - Success strategies
        
        Format the response in a clear, structured way with sections and bullet points.
        Make the response realistic and specific to {company_name}'s known practices.
        """
        
        try:
            response = self._call_llm_api(prompt)
            logging.info({"From":["Start"],"Desc": f"""{response}""","To":"CompanyResearcher"})
            return response
        except Exception as e:
            return f"Error generating company insights: {str(e)}"

    def get_leetcode_recommendations(self, job_description):
        """Generate LeetCode problem recommendations based on job requirements."""
        prompt = f"""
        Based on the following job description, recommend LeetCode problems and topics to prepare for technical interviews:
        
        Job Description:
        {job_description}
        
        Please provide:
        1. Core Topics to Focus On:
           - List key data structures
           - Important algorithms
           - System design concepts (if applicable)
        
        2. Recommended LeetCode Problems:
           - Easy level problems (5-7 problems)
           - Medium level problems (7-10 problems)
           - Hard level problems (3-5 problems)
        
        3. Study Plan:
           - Week-by-week preparation guide
           - Time allocation suggestions
           - Practice strategies
        
        4. Interview Success Tips:
           - Problem-solving approach
           - Time management during interviews
           - Common patterns to recognize
        
        Format the response in a clear, structured way with sections and bullet points.
        Include specific LeetCode problem numbers and names where possible.
        """
        
        try:
            response = self._call_llm_api(prompt)
            logging.info({"From":["Start"],"Desc": f"""{response}""","To":"LeetcodeResearcher"})
            return response
        except Exception as e:
            return f"Error generating LeetCode recommendations: {str(e)}"

    def generate_interview_prep(self, resume_text, job_description, company_name):
        """Generate interview preparation content based on resume and job description."""
        prompt = f"""
        Based on the following information, provide comprehensive interview preparation guidance:
        
        Resume:
        {resume_text}
        
        Job Description:
        {job_description}
        
        Company:
        {company_name}
        
        Please provide:
        1. Technical Interview Preparation:
           - Key technical concepts to review
           - Sample technical questions
           - Coding challenges to practice
        
        2. Behavioral Interview Questions:
           - Role-specific scenarios
           - Leadership examples to prepare
           - Problem-solving stories to develop
        
        3. Company-Specific Preparation:
           - Company values alignment
           - Recent company developments
           - Industry knowledge to demonstrate
        
        4. Interview Strategy:
           - How to present your experience
           - Answering difficult questions
           - Questions to ask the interviewer
        
        5. Success Tips:
           - Pre-interview preparation
           - Day-of interview recommendations
           - Follow-up strategies
        
        Format the response in a clear, structured way with sections and bullet points.
        Make recommendations specific to the role and company where possible.
        """
        
        try:
            response = self._call_llm_api(prompt)
            logging.info({"From":["Start"],"Desc": f"""{response}""","To":"InterviewPrepResearcher"})
            return response
        except Exception as e:
            return f"Error generating interview preparation: {str(e)}"

# Example usage
if __name__ == "__main__":
    analyzer = ResumeAnalyzer()
    with open("resume.pdf", "rb") as file:
        resume_text = analyzer.extract_text_from_file(file)
        analysis = analyzer.analyze_resume(resume_text)
        print(analysis) 